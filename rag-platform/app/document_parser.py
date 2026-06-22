"""
document_parser.py
Extracts plain text from PDF, DOCX, and TXT files.
All parsing is done locally — no external API calls.
"""

import os
from pypdf import PdfReader
from docx import Document as DocxDocument


def parse_pdf(file_path: str) -> str:
    """Extract text from a PDF file, page by page."""
    text_chunks = []
    reader = PdfReader(file_path)
    for page_num, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        if page_text.strip():
            text_chunks.append(f"[Page {page_num}]\n{page_text}")
    return "\n\n".join(text_chunks)


def parse_docx(file_path: str) -> str:
    """Extract text from a Word document, including tables."""
    doc = DocxDocument(file_path)
    text_chunks = []

    for para in doc.paragraphs:
        if para.text.strip():
            text_chunks.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                text_chunks.append(row_text)

    return "\n".join(text_chunks)


def parse_txt(file_path: str) -> str:
    """Read a plain text file."""
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def parse_document(file_path: str) -> str:
    """
    Dispatch to the correct parser based on file extension.
    Returns extracted plain text, or raises ValueError for unsupported types.
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return parse_pdf(file_path)
    elif ext == ".docx":
        return parse_docx(file_path)
    elif ext == ".txt":
        return parse_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Supported: .pdf, .docx, .txt")
